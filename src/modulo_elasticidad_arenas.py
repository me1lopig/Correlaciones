import streamlit as st
from docx import Document
import io
import pandas as pd
import math

def calcular_modulo_elasticidad(Nspt=None):
    """
    Calcula el módulo de elasticidad (E) según los datos disponibles.
    Solo aplica fórmulas para las que todos los parámetros necesarios estén disponibles.
    """
    resultados = {}
    formulas_usadas = {}

    def valor_aceptable(valor):
        return valor if valor is not None and valor >= 0 else None

    Nspt = valor_aceptable(Nspt)

    if Nspt is not None:
        # Fórmula de Webb, 1974 para Arenas arcillosas
        E = 0.324 * (Nspt + 15)
        resultados['Webb, 1974 (Arenas arcillosas)'] = E
        formulas_usadas['Webb, 1974 (Arenas arcillosas)'] = {'formula': 'E = 0.324 × (Nspt + 15)', 'parametros': ['Nspt']}

        # Fórmula de Webb, 1974 para Casos intermedios
        E = 0.392 * (Nspt + 12)
        resultados['Webb, 1974 (Casos intermedios)'] = E
        formulas_usadas['Webb, 1974 (Casos intermedios)'] = {'formula': 'E = 0.392 × (Nspt + 12)', 'parametros': ['Nspt']}

        # Fórmula de Denver, 1982 para Arenas
        E = 7 * math.sqrt(Nspt)
        resultados['Denver, 1982 (Arenas)'] = E
        formulas_usadas['Denver, 1982 (Arenas)'] = {'formula': 'E = 7 × √Nspt', 'parametros': ['Nspt']}

        # Fórmula de Meigh y Nixon, 1961 para Limos y limos arenosos
        E = 0.491 * Nspt
        resultados['Meigh y Nixon, 1961 (Limos y limos arenosos)'] = E
        formulas_usadas['Meigh y Nixon, 1961 (Limos y limos arenosos)'] = {'formula': 'E = 0.491 × Nspt', 'parametros': ['Nspt']}

        # Fórmula de Meigh y Nixon, 1961 para Arenas finas
        E = 0.785 * Nspt
        resultados['Meigh y Nixon, 1961 (Arenas finas)'] = E
        formulas_usadas['Meigh y Nixon, 1961 (Arenas finas)'] = {'formula': 'E = 0.785 × Nspt', 'parametros': ['Nspt']}

        # Fórmula de Wroth y Nowatzki, 1986 para Gravas
        E = 2.22 * Nspt ** 0.886
        resultados['Wroth y Nowatzki, 1986 (Gravas)'] = E
        formulas_usadas['Wroth y Nowatzki, 1986 (Gravas)'] = {'formula': 'E = 2.22 × Nspt^0.886', 'parametros': ['Nspt']}

        # Fórmula de Bowles, 1996 para Arenas normalmente consolidadas
        E = 0.491 * (Nspt + 15)
        resultados['Bowles, 1996 (Arenas normalmente consolidadas)'] = E
        formulas_usadas['Bowles, 1996 (Arenas normalmente consolidadas)'] = {'formula': 'E = 0.491 × (Nspt + 15)', 'parametros': ['Nspt']}

        # Fórmula de Bowles, 1996 para Arenas en general
        E = 0.491 * (Nspt + 15)
        resultados['Bowles, 1996 (Arenas en general)'] = E
        formulas_usadas['Bowles, 1996 (Arenas en general)'] = {'formula': 'E = 0.491 × (Nspt + 15)', 'parametros': ['Nspt']}

        # Fórmula de Bowles, 1996 para Gravas (N ≤ 15)
        E = 0.589 * (Nspt + 6)
        resultados['Bowles, 1996 (Gravas, N ≤ 15)'] = E
        formulas_usadas['Bowles, 1996 (Gravas, N ≤ 15)'] = {'formula': 'E = 0.589 × (Nspt + 6)', 'parametros': ['Nspt']}

        # Fórmula de Bowles, 1996 para Gravas (N > 15)
        E = 0.589 * (Nspt + 6) + 1.962
        resultados['Bowles, 1996 (Gravas, N > 15)'] = E
        formulas_usadas['Bowles, 1996 (Gravas, N > 15)'] = {'formula': 'E = 0.589 × (Nspt + 6) + 1.962', 'parametros': ['Nspt']}

        # Fórmula de Bowles, 1996 para Gravas y arenas (N ≤ 15)
        E = 1.177 * (Nspt + 6)
        resultados['Bowles, 1996 (Gravas y arenas, N ≤ 15)'] = E
        formulas_usadas['Bowles, 1996 (Gravas y arenas, N ≤ 15)'] = {'formula': 'E = 1.177 × (Nspt + 6)', 'parametros': ['Nspt']}

        # Fórmula de Bowles, 1996 para Gravas y arenas (N > 15)
        E = 1.177 * (Nspt + 6) + 3.924
        resultados['Bowles, 1996 (Gravas y arenas, N > 15)'] = E
        formulas_usadas['Bowles, 1996 (Gravas y arenas, N > 15)'] = {'formula': 'E = 1.177 × (Nspt + 6) + 3.924', 'parametros': ['Nspt']}

        # Fórmula de Schertmann, 1970 para Arenas
        E = 0.785 * Nspt
        resultados['Schertmann, 1970 (Arenas)'] = E
        formulas_usadas['Schertmann, 1970 (Arenas)'] = {'formula': 'E = 0.785 × Nspt', 'parametros': ['Nspt']}

        # Fórmula de D’Polonia y otros para Arenas normalmente consolidadas
        E = 21.09 + 1.04 * Nspt
        resultados['D’Polonia y otros (Arenas normalmente consolidadas)'] = E
        formulas_usadas['D’Polonia y otros (Arenas normalmente consolidadas)'] = {'formula': 'E = 21.09 + 1.04 × Nspt', 'parametros': ['Nspt']}

        # Fórmula de D’Polonia y otros para Arenas preconsolidadas
        E = 52.97 + 1.32 * Nspt
        resultados['D’Polonia y otros (Arenas preconsolidadas)'] = E
        formulas_usadas['D’Polonia y otros (Arenas preconsolidadas)'] = {'formula': 'E = 52.97 + 1.32 × Nspt', 'parametros': ['Nspt']}

    return resultados, formulas_usadas

def generar_informe(Nspt, resultados, formulas_usadas):
    doc = Document()
    doc.add_heading('Informe de Cálculo del Módulo de Elasticidad (E)', level=1)

    doc.add_heading('Datos Introducidos', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells

    hdr_cells[0].text = 'Parámetro'
    hdr_cells[1].text = 'Valor'
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    datos = []
    if Nspt is not None:
        datos.append(("Número de golpes SPT (Nspt)", Nspt))

    for parametro, valor in datos:
        row_cells = table.add_row().cells
        row_cells[0].text = parametro
        row_cells[1].text = str(valor)

    doc.add_heading('Resultados', level=2)
    doc.add_paragraph(f"Se calcularon {len(resultados)} correlaciones para E con los datos proporcionados.")

    for metodo, valor in resultados.items():
        doc.add_heading(metodo, level=3)
        formula_info = formulas_usadas[metodo]
        p = doc.add_paragraph()
        run = p.add_run(f"Fórmula: {formula_info['formula']}")
        run.bold = True
        doc.add_paragraph(f"Aplicación: {metodo.split('(')[1].rstrip(')')}")
        doc.add_paragraph(f"Parámetros usados: {', '.join(formula_info['parametros'])}")
        doc.add_paragraph(f"Resultado: E ≈ {valor:.2f} MPa")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def main():
    st.set_page_config(layout="wide")
    st.title("Calculadora del Módulo de Elasticidad (E) para Arenas")
    st.markdown("Introduce el número de golpes SPT para calcular E según diferentes correlaciones.")

    if 'data' not in st.session_state:
        st.session_state.data = {
            "Símbolo": ["Nspt"],
            "Descripción del Parámetro": ["Número de golpes SPT"],
            "Valor": [None],
            "Unidad": ["-"]
        }

    col1, col2 = st.columns([1, 1])

    with col1:
        df = pd.DataFrame(st.session_state.data)
        edited_df = st.data_editor(
            df,
            num_rows="fixed",
            hide_index=True,
            column_config={
                "Símbolo": st.column_config.TextColumn("Símbolo", disabled=True),
                "Descripción del Parámetro": st.column_config.TextColumn("Descripción del Parámetro", disabled=True),
                "Unidad": st.column_config.TextColumn("Unidad", disabled=True)
            },
            key="data_editor"
        )

        if st.button("Calcular E"):
            try:
                Nspt = float(edited_df.loc[edited_df["Símbolo"] == "Nspt", "Valor"].values[0]) if edited_df.loc[edited_df["Símbolo"] == "Nspt", "Valor"].values[0] and float(edited_df.loc[edited_df["Símbolo"] == "Nspt", "Valor"].values[0]) >= 0 else None
            except (ValueError, TypeError):
                Nspt = None

            resultados, formulas_usadas = calcular_modulo_elasticidad(Nspt=Nspt)

            if resultados:
                st.subheader("Resultados de E según correlaciones aplicables:")
                for metodo, valor in resultados.items():
                    st.markdown(f"**{metodo}**")
                    st.write(f"Fórmula: {formulas_usadas[metodo]['formula']}")
                    st.write(f"Aplicación: {metodo.split('(')[1].rstrip(')')}")
                    st.write(f"Parámetros usados: {', '.join(formulas_usadas[metodo]['parametros'])}")
                    st.write(f"Resultado: E ≈ {valor:.2f} MPa")
                    st.markdown("---")

                informe_buffer = generar_informe(Nspt, resultados, formulas_usadas)
                st.download_button(
                    label="Descargar informe en Word",
                    data=informe_buffer,
                    file_name="informe_modulo_elasticidad.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.warning("No hay suficientes datos para aplicar ninguna correlación.")

    with col2:
        with st.expander("Correlaciones Disponibles", expanded=False):
            st.markdown("""
            <style>
            .formula {
                border-left: 3px solid #4CAF50;
                padding-left: 10px;
                margin-bottom: 15px;
            }
            </style>
            """, unsafe_allow_html=True)

            correlaciones = [
                {"name": "Webb, 1974 (Arenas arcillosas)", "formula": "E = 0.324 × (Nspt + 15)", "params": "Nspt", "aplicacion": "Arenas arcillosas"},
                {"name": "Webb, 1974 (Casos intermedios)", "formula": "E = 0.392 × (Nspt + 12)", "params": "Nspt", "aplicacion": "Casos intermedios"},
                {"name": "Denver, 1982 (Arenas)", "formula": "E = 7 × √Nspt", "params": "Nspt", "aplicacion": "Arenas"},
                {"name": "Meigh y Nixon, 1961 (Limos y limos arenosos)", "formula": "E = 0.491 × Nspt", "params": "Nspt", "aplicacion": "Limos y limos arenosos"},
                {"name": "Meigh y Nixon, 1961 (Arenas finas)", "formula": "E = 0.785 × Nspt", "params": "Nspt", "aplicacion": "Arenas finas"},
                {"name": "Wroth y Nowatzki, 1986 (Gravas)", "formula": "E = 2.22 × Nspt^0.886", "params": "Nspt", "aplicacion": "Gravas"},
                {"name": "Bowles, 1996 (Arenas normalmente consolidadas)", "formula": "E = 0.491 × (Nspt + 15)", "params": "Nspt", "aplicacion": "Arenas normalmente consolidadas"},
                {"name": "Bowles, 1996 (Arenas en general)", "formula": "E = 0.491 × (Nspt + 15)", "params": "Nspt", "aplicacion": "Arenas en general"},
                {"name": "Bowles, 1996 (Gravas, N ≤ 15)", "formula": "E = 0.589 × (Nspt + 6)", "params": "Nspt", "aplicacion": "Gravas (N ≤ 15)"},
                {"name": "Bowles, 1996 (Gravas, N > 15)", "formula": "E = 0.589 × (Nspt + 6) + 1.962", "params": "Nspt", "aplicacion": "Gravas (N > 15)"},
                {"name": "Bowles, 1996 (Gravas y arenas, N ≤ 15)", "formula": "E = 1.177 × (Nspt + 6)", "params": "Nspt", "aplicacion": "Gravas y arenas (N ≤ 15)"},
                {"name": "Bowles, 1996 (Gravas y arenas, N > 15)", "formula": "E = 1.177 × (Nspt + 6) + 3.924", "params": "Nspt", "aplicacion": "Gravas y arenas (N > 15)"},
                {"name": "Schertmann, 1970 (Arenas)", "formula": "E = 0.785 × Nspt", "params": "Nspt", "aplicacion": "Arenas"},
                {"name": "D’Polonia y otros (Arenas normalmente consolidadas)", "formula": "E = 21.09 + 1.04 × Nspt", "params": "Nspt", "aplicacion": "Arenas normalmente consolidadas"},
                {"name": "D’Polonia y otros (Arenas preconsolidadas)", "formula": "E = 52.97 + 1.32 × Nspt", "params": "Nspt", "aplicacion": "Arenas preconsolidadas"}
            ]

            for correlacion in correlaciones:
                st.markdown(f'<div class="formula"><b>{correlacion["name"]}</b><br>Fórmula: {correlacion["formula"]}<br>Aplicación: {correlacion["aplicacion"]}<br>Parámetros: {correlacion["params"]}</div>', unsafe_allow_html=True)


    

if __name__ == "__main__":
    main()
