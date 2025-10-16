import streamlit as st
from docx import Document
import io
import pandas as pd

def calcular_angulo_rozamiento(IP=None, Nspt=None):
    """
    Calcula el ángulo de rozamiento (φ) según los datos disponibles.
    Solo aplica fórmulas para las que todos los parámetros necesarios estén disponibles.
    """
    resultados = {}
    formulas_usadas = {}

    def valor_aceptable(valor):
        return valor if valor is not None and valor > 0 else None

    IP = valor_aceptable(IP)
    Nspt = valor_aceptable(Nspt)

    if IP is not None:
        # Fórmula de Jimenes Salas y Justo Alpañes
        phi = 1.1616 * IP
        resultados['Jimenes Salas y Justo Alpañes'] = phi if phi > 0 else "Valor no aceptable para los datos introducidos"
        formulas_usadas['Jimenes Salas y Justo Alpañes'] = {'formula': 'φ ≈ 1.1616 × IP', 'parametros': ['IP']}

    if Nspt is not None:
        # Fórmula de Peck
        phi = 2.0986 * Nspt
        resultados['Peck'] = phi if phi <= 45 else 45  # Limitar a 45 grados
        formulas_usadas['Peck'] = {'formula': 'φ ≈ 2.0986 × Nspt', 'parametros': ['Nspt']}

        # Fórmula de Muromachi, 1974
        phi = 2.2370 * Nspt
        resultados['Muromachi, 1974'] = phi if phi <= 45 else 45
        formulas_usadas['Muromachi, 1974'] = {'formula': 'φ ≈ 2.2370 × Nspt', 'parametros': ['Nspt']}

        # Fórmula de Terzaghi & Peck 1948
        phi = 2.1500 * Nspt
        resultados['Terzaghi & Peck 1948'] = phi if phi <= 45 else 45
        formulas_usadas['Terzaghi & Peck 1948'] = {'formula': 'φ ≈ 2.1500 × Nspt', 'parametros': ['Nspt']}

        # Fórmula de Kishida, 1969
        phi = 2.1547 * Nspt
        resultados['Kishida, 1969'] = phi if phi <= 45 else 45
        formulas_usadas['Kishida, 1969'] = {'formula': 'φ ≈ 2.1547 × Nspt', 'parametros': ['Nspt']}

        # Fórmula de Japan National Railway, 1999
        phi = 2.1000 * Nspt
        resultados['Japan National Railway, 1999'] = phi if phi <= 45 else 45
        formulas_usadas['Japan National Railway, 1999'] = {'formula': 'φ ≈ 2.1000 × Nspt', 'parametros': ['Nspt']}

        # Fórmula de Japan Road Bureau, 1986
        phi = 1.7906 * Nspt + 15
        resultados['Japan Road Bureau, 1986'] = phi if phi <= 45 else 45
        formulas_usadas['Japan Road Bureau, 1986'] = {'formula': 'φ ≈ 1.7906 × Nspt + 15', 'parametros': ['Nspt']}

        # Fórmula de Hatanaka & Uchida, 1996
        phi = 2.4880 * Nspt
        resultados['Hatanaka & Uchida, 1996'] = phi if phi <= 45 else 45
        formulas_usadas['Hatanaka & Uchida, 1996'] = {'formula': 'φ ≈ 2.4880 × Nspt', 'parametros': ['Nspt']}

        # Fórmula de Montenegro & Gonzalez
        phi = 2.1657 * Nspt
        resultados['Montenegro & Gonzalez'] = phi if phi <= 45 else 45
        formulas_usadas['Montenegro & Gonzalez'] = {'formula': 'φ ≈ 2.1657 × Nspt', 'parametros': ['Nspt']}

        # Fórmula de Shiol-Fukuni 1982
        phi = 2.0000 * Nspt + 15
        resultados['Shiol-Fukuni 1982'] = phi if phi <= 45 else 45
        formulas_usadas['Shiol-Fukuni 1982'] = {'formula': 'φ ≈ 2.0000 × Nspt + 15', 'parametros': ['Nspt']}

    return resultados, formulas_usadas

def generar_informe(IP, Nspt, resultados, formulas_usadas):
    doc = Document()
    doc.add_heading('Informe de Cálculo del Ángulo de Rozamiento (φ)', level=1)

    doc.add_heading('Datos Introducidos', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells

    hdr_cells[0].text = 'Parámetro'
    hdr_cells[1].text = 'Valor'
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    datos = []
    if IP is not None:
        datos.append(("Índice plástico (IP, %)", IP))
    if Nspt is not None:
        datos.append(("Número de golpes SPT (Nspt)", Nspt))

    for parametro, valor in datos:
        row_cells = table.add_row().cells
        row_cells[0].text = parametro
        row_cells[1].text = str(valor)

    doc.add_heading('Resultados', level=2)
    doc.add_paragraph(f"Se calcularon {len(resultados)} correlaciones para φ con los datos proporcionados.")

    for metodo, valor in resultados.items():
        doc.add_heading(metodo, level=3)
        formula_info = formulas_usadas[metodo]
        p = doc.add_paragraph()
        run = p.add_run(f"Fórmula: {formula_info['formula']}")
        run.bold = True
        doc.add_paragraph(f"Parámetros usados: {', '.join(formula_info['parametros'])}")
        if isinstance(valor, str):
            doc.add_paragraph(f"Resultado: {valor}")
        else:
            doc.add_paragraph(f"Resultado: φ ≈ {valor:.2f}°")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def main():
    st.set_page_config(layout="wide")
    st.title("Calculadora del Ángulo de Rozamiento (φ)")
    st.markdown("Introduce los datos disponibles para calcular φ según diferentes correlaciones.")

    if 'data' not in st.session_state:
        st.session_state.data = {
            "Símbolo": ["IP", "Nspt"],
            "Descripción del Parámetro": [
                "Índice plástico", "Número de golpes SPT"
            ],
            "Valor": [None, None],
            "Unidad": ["%", "-"]
        }

    col1, col2 = st.columns([1, 1])

    with col1:
        df = pd.DataFrame(st.session_state.data)
        edited_df = st.data_editor(
            df,
            num_rows="fixed",
            hide_index=True,
            column_config={
                "Símbolo": st.column_config.TextColumn("Símbolo", disabled=True),
                "Descripción del Parámetro": st.column_config.TextColumn("Descripción del Parámetro", disabled=True),
                "Unidad": st.column_config.TextColumn("Unidad", disabled=True)
            },
            key="data_editor"
        )

        if st.button("Calcular φ"):
            try:
                IP = float(edited_df.loc[edited_df["Símbolo"] == "IP", "Valor"].values[0]) if edited_df.loc[edited_df["Símbolo"] == "IP", "Valor"].values[0] and float(edited_df.loc[edited_df["Símbolo"] == "IP", "Valor"].values[0]) > 0 else None
            except (ValueError, TypeError):
                IP = None

            try:
                Nspt = float(edited_df.loc[edited_df["Símbolo"] == "Nspt", "Valor"].values[0]) if edited_df.loc[edited_df["Símbolo"] == "Nspt", "Valor"].values[0] and float(edited_df.loc[edited_df["Símbolo"] == "Nspt", "Valor"].values[0]) > 0 else None
            except (ValueError, TypeError):
                Nspt = None

            resultados, formulas_usadas = calcular_angulo_rozamiento(IP=IP, Nspt=Nspt)

            if resultados:
                st.subheader("Resultados de φ según correlaciones aplicables:")
                for metodo, valor in resultados.items():
                    st.markdown(f"**{metodo}**")
                    st.write(f"Fórmula: {formulas_usadas[metodo]['formula']}")
                    st.write(f"Parámetros usados: {', '.join(formulas_usadas[metodo]['parametros'])}")
                    if isinstance(valor, str):
                        st.write(f"Resultado: {valor}")
                    else:
                        st.write(f"Resultado: φ ≈ {valor:.2f}°")
                    st.markdown("---")

                informe_buffer = generar_informe(IP, Nspt, resultados, formulas_usadas)
                st.download_button(
                    label="Descargar informe en Word",
                    data=informe_buffer,
                    file_name="informe_angulo_rozamiento.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.warning("No hay suficientes datos para aplicar ninguna correlación.")

    with col2:
        with st.expander("Correlaciones Disponibles", expanded=False):
            st.markdown("""
            <style>
            .formula {
                border-left: 3px solid #4CAF50;
                padding-left: 10px;
                margin-bottom: 15px;
            }
            </style>
            """, unsafe_allow_html=True)

            correlaciones = [
                {"name": "Jimenes Salas y Justo Alpañes", "formula": "φ ≈ 1.1616 × IP", "params": "IP"},
                {"name": "Peck", "formula": "φ ≈ 2.0986 × Nspt", "params": "Nspt"},
                {"name": "Muromachi, 1974", "formula": "φ ≈ 2.2370 × Nspt", "params": "Nspt"},
                {"name": "Terzaghi & Peck 1948", "formula": "φ ≈ 2.1500 × Nspt", "params": "Nspt"},
                {"name": "Kishida, 1969", "formula": "φ ≈ 2.1547 × Nspt", "params": "Nspt"},
                {"name": "Japan National Railway, 1999", "formula": "φ ≈ 2.1000 × Nspt", "params": "Nspt"},
                {"name": "Japan Road Bureau, 1986", "formula": "φ ≈ 1.7906 × Nspt + 15", "params": "Nspt"},
                {"name": "Hatanaka & Uchida, 1996", "formula": "φ ≈ 2.4880 × Nspt", "params": "Nspt"},
                {"name": "Montenegro & Gonzalez", "formula": "φ ≈ 2.1657 × Nspt", "params": "Nspt"},
                {"name": "Shiol-Fukuni 1982", "formula": "φ ≈ 2.0000 × Nspt + 15", "params": "Nspt"}
            ]

            for correlacion in correlaciones:
                st.markdown(f'<div class="formula"><b>{correlacion["name"]}</b><br>Fórmula: {correlacion["formula"]}<br>Parámetros: {correlacion["params"]}</div>', unsafe_allow_html=True)

if __name__ == "__main__":
    main()
