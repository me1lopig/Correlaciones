import streamlit as st
import pandas as pd
import plotly.express as px
from io import BytesIO
import copy
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

# ==============================================================================
# CONFIGURACIÓN Y CONSTANTES
# ==============================================================================
st.set_page_config(page_title="E Suelos Granulares", layout="wide")

FACTOR_KG_CM2_A_MPA = 0.0980665

BIBLIOGRAFIA = [
    "Begemann, H. K. S. (1974). General report for central and western Europe. Proceedings of the ESOPT, Stockholm.",
    "Bowles, J. E. (1996). Foundation Analysis and Design (5th ed.). McGraw-Hill.",
    "D'Appolonia, D. J., D'Appolonia, E., & Brissette, R. F. (1970). Settlement of spread footings on sand. Journal of the Soil Mechanics and Foundations Division, ASCE.",
    "Denver, H. (1982). Modulus of elasticity for sand determined by SPT and CPT. Proceedings of the 2nd ESOPT, Amsterdam.",
    "Meigh, A. C., & Nixon, I. K. (1961). Comparison of in-situ tests for granular soils. Proceedings of the 5th ICSMFE, Paris.",
    "Schmertmann, J. H. (1970). Static cone to compute static settlement over sand. Journal of the Soil Mechanics and Foundations Division, ASCE.",
    "Webb, D. L. (1969). Settlement of structures on deep alluvial sandy sediments. Proceedings of the Conference on In Situ Investigations in Soils and Rocks, BGS, London.",
    "Wrench, B. P., & Nowatzki, E. A. (1986). A relationship between deformation modulus and SPT N for gravels. The Civil Engineer in South Africa."
]

# ==========================================
# 1. LÓGICA MATEMÁTICA
# ==========================================
def calcular_datos_base(N_spt: int):
    """Genera todos los cálculos posibles."""
    resultados = []
    
    # Webb (1969)
    resultados.append({"Autor": "Webb (1969)", "Aplicación": "Arenas arcillosas", "Fórmula Original": "3.3·(N+15) kg/cm²", "E (MPa)": 3.3 * (N_spt + 15) * FACTOR_KG_CM2_A_MPA})
    resultados.append({"Autor": "Webb (1969)", "Aplicación": "Suelos intermedios", "Fórmula Original": "4·(N+12) kg/cm²", "E (MPa)": 4.0 * (N_spt + 12) * FACTOR_KG_CM2_A_MPA})

    # Meigh & Nixon (1961)
    resultados.append({"Autor": "Meigh & Nixon (1961)", "Aplicación": "Limos y limos arenosos", "Fórmula Original": "5·N kg/cm²", "E (MPa)": 5.0 * N_spt * FACTOR_KG_CM2_A_MPA})
    resultados.append({"Autor": "Meigh & Nixon (1961)", "Aplicación": "Arenas finas", "Fórmula Original": "8·N kg/cm²", "E (MPa)": 8.0 * N_spt * FACTOR_KG_CM2_A_MPA})

    # Bowles (1996)
    resultados.append({"Autor": "Bowles (1996)", "Aplicación": "Arenas (NC)", "Fórmula Original": "5·(N+15) kg/cm²", "E (MPa)": 5.0 * (N_spt + 15) * FACTOR_KG_CM2_A_MPA})
    
    if N_spt <= 15: val, txt = 6.0 * (N_spt + 6) * FACTOR_KG_CM2_A_MPA, "6·(N+6) kg/cm²"
    else: val, txt = (6.0 * (N_spt + 6) + 20.0) * FACTOR_KG_CM2_A_MPA, "6·(N+6) + 20 kg/cm²"
    resultados.append({"Autor": "Bowles (1996)", "Aplicación": "Gravas", "Fórmula Original": txt, "E (MPa)": val})

    # Begemann (1974)
    if N_spt <= 15: val, txt = 12.0 * (N_spt + 6) * FACTOR_KG_CM2_A_MPA, "12·(N+6) kg/cm²"
    else: val, txt = (12.0 * (N_spt + 6) + 40.0) * FACTOR_KG_CM2_A_MPA, "12·(N+6) + 40 kg/cm²"
    resultados.append({"Autor": "Begemann (1974)", "Aplicación": "Gravas y Arenas", "Fórmula Original": txt, "E (MPa)": val})

    # Schmertmann (1970)
    resultados.append({"Autor": "Schmertmann (1970)", "Aplicación": "Arenas", "Fórmula Original": "8·N kg/cm²", "E (MPa)": 8.0 * N_spt * FACTOR_KG_CM2_A_MPA})
    
    # D'Appolonia (1970)
    resultados.append({"Autor": "D'Appolonia (1970)", "Aplicación": "Arenas (NC)", "Fórmula Original": "215 + 10.6·N kg/cm²", "E (MPa)": (215 + 10.6 * N_spt) * FACTOR_KG_CM2_A_MPA})
    resultados.append({"Autor": "D'Appolonia (1970)", "Aplicación": "Arenas (Precons.)", "Fórmula Original": "540 + 13.5·N kg/cm²", "E (MPa)": (540 + 13.5 * N_spt) * FACTOR_KG_CM2_A_MPA})

    # Denver (1982)
    resultados.append({"Autor": "Denver (1982)", "Aplicación": "Arenas (General)", "Fórmula Original": "7·√N (MPa)", "E (MPa)": 7 * (N_spt ** 0.5)})

    # Wrench & Nowatzki (1986)
    resultados.append({"Autor": "Wrench & Nowatzki (1986)", "Aplicación": "Gravas", "Fórmula Original": "2.22·N^0.888 (MPa)", "E (MPa)": 2.22 * (N_spt ** 0.888)})

    return pd.DataFrame(resultados)

# ==========================================
# 2. FILTRADO LÓGICO
# ==========================================
def aplicar_filtro_suelo(df, tipo_suelo):
    if tipo_suelo == "Mostrar Todo":
        return df
    
    df["_temp_app"] = df["Aplicación"].str.lower()
    
    if tipo_suelo == "Arenas":
        mask = df["_temp_app"].str.contains("arena")
    elif tipo_suelo == "Gravas":
        mask = df["_temp_app"].str.contains("grava")
    elif tipo_suelo == "Limos":
        mask = df["_temp_app"].str.contains("limo")
    elif tipo_suelo == "Suelos Intermedios":
        mask = df["_temp_app"].str.contains("intermedios")
    else:
        mask = [True] * len(df)
        
    df_filtrado = df[mask].drop(columns=["_temp_app"])
    return df_filtrado

# ==========================================
# 3. GENERADOR DE INFORME WORD (Estilo CTE_2219)
# ==========================================
def generar_docx(n_val, df_final, df_stats, fig_plotly_original, tipo_suelo_selec):
    doc = Document()
    
    # --- 1. CONFIGURACIÓN DE ESTILOS DEL DOCUMENTO MODELO ---
    # Colores extraídos del documento modelo
    COLOR_TITULO = RGBColor(0x17, 0x36, 0x5D)  # Azul Oscuro
    COLOR_HEADING = RGBColor(0x36, 0x5F, 0x91) # Azul Medio
    
    # Configurar Márgenes (1.25" laterales, 1.0" verticales)
    for section in doc.sections:
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # Configurar Fuente Normal (Calibri 11pt)
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)

    # Configurar Título (Title)
    style_title = doc.styles['Title']
    style_title.font.name = 'Calibri Light' # Suele ser Light en temas modernos de Word
    style_title.font.size = Pt(26)
    style_title.font.color.rgb = COLOR_TITULO
    
    # Configurar Encabezados (Heading 1)
    style_h1 = doc.styles['Heading 1']
    style_h1.font.name = 'Calibri Light'
    style_h1.font.size = Pt(14)
    style_h1.font.color.rgb = COLOR_HEADING
    
    # --- 2. CONTENIDO DEL INFORME ---

    # Título Principal
    doc.add_heading('Informe estimación Módulo de Elasticidad en Suelos Granulares', 0).alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Fecha negrita
    p_fecha = doc.add_paragraph()
    p_fecha.add_run(f'Fecha de emisión: {pd.Timestamp.now().strftime("%d/%m/%Y")}').bold = True
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.LEFT
    #doc.add_paragraph('---')

    # Sección 1
    doc.add_heading('1. Datos de Entrada', level=1)
    p = doc.add_paragraph()
    p.add_run(f'• Valor N (SPT) de diseño: ').bold = True
    p.add_run(f'{n_val} golpes/30 cm')
    
    p2 = doc.add_paragraph()
    p2.add_run(f'• Tipo de suelo seleccionado: ').bold = True
    p2.add_run(f'{tipo_suelo_selec}')

    # Sección 2 - Tabla Principal
    doc.add_heading('2. Métodos de Cálculo Seleccionados', level=1)
    
    # Tabla con estilo "Light List Accent 1" (Estilo del modelo)
    table = doc.add_table(rows=1, cols=4)
    try:
        table.style = 'Light List Accent 1'
    except:
        table.style = 'Table Grid' # Fallback si no encuentra el estilo exacto
        
    table.autofit = False 
    table.allow_autofit = False
    
    widths = [Inches(2.0), Inches(1.5), Inches(2.0), Inches(1.0)] # Ajustados para margen 1.25"
    headers = ['Autor', 'Aplicación', 'Fórmula Original', 'E (MPa)']
    
    # Encabezados
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.width = widths[i]
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
    # Filas de datos
    for _, row in df_final.iterrows():
        row_cells = table.add_row().cells
        textos = [str(row['Autor']), str(row['Aplicación']), str(row['Fórmula Original']), f"{row['E (MPa)']:.2f}"]
        for idx, txt in enumerate(textos):
            row_cells[idx].text = txt
            row_cells[idx].width = widths[idx]
            row_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            row_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sección 3 - Estadísticas
    doc.add_heading('3. Análisis Estadístico', level=1)
    if df_stats is not None:
        doc.add_paragraph('Resumen estadístico de los métodos seleccionados:')
        stat_table = doc.add_table(rows=1, cols=5)
        try:
            stat_table.style = 'Light List Accent 1'
        except:
            stat_table.style = 'Table Grid'
            
        stat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        stat_table.autofit = True
        
        headers_stat = ['Mínimo', 'Máximo', 'Promedio', 'Mediana', 'Desv. Típica']
        for i, h in enumerate(headers_stat):
            cell = stat_table.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT 
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(10)

        vals = stat_table.add_row().cells
        data_vals = [f"{df_stats.iloc[0][k]:.2f}" for k in headers_stat]
        for i, val in enumerate(data_vals):
            vals[i].text = val
            vals[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT 
            vals[i].paragraphs[0].runs[0].font.size = Pt(10)
    else:
        doc.add_paragraph('No procede cálculo estadístico (selección insuficiente).')

    # Sección 4 - Gráfica
    doc.add_heading('4. Gráfica Comparativa', level=1)
    try:
        # Clonamos y re-configuramos
        fig_word = copy.deepcopy(fig_plotly_original)
        
        fig_word.update_traces(
            textfont_size=14,
            textfont_color='black'
        )
        fig_word.update_layout(
            uniformtext_minsize=14, 
            uniformtext_mode='show'
        )

        # Exportamos con alta resolución
        img_bytes = fig_word.to_image(format="png", width=1300, height=len(df_final)*60 + 200, scale=3)
        
        doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Ajustamos el ancho al nuevo ancho de página (6.0 pulgadas aprox con margen 1.25)
        doc.paragraphs[-1].add_run().add_picture(BytesIO(img_bytes), width=Inches(6.0))
    except Exception as e:
        doc.add_paragraph(f"[Gráfica no disponible: {str(e)}]")

    doc.add_heading('5. Referencias Bibliográficas', level=1)
    for ref in BIBLIOGRAFIA:
        #p = doc.add_paragraph(ref)
        p = doc.add_paragraph(ref, style='List Bullet')
        p.style.font.name = 'Calibri'
        p.style.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(6)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# 4. INTERFAZ STREAMLIT
# ==========================================

st.title("🏖️ Estimación Módulo de Elasticidad en Suelos Granulares")
st.markdown("---")

# --- CONTROL DE ESTADO ---
if 'n_previo' not in st.session_state: st.session_state.n_previo = 0
if 'tipo_suelo_previo' not in st.session_state: st.session_state.tipo_suelo_previo = None 
if 'selecciones' not in st.session_state: st.session_state.selecciones = {}

with st.sidebar:
    st.header("1.⚙️Valores de N (SPT)")
    n_spt = st.number_input("Valor N (SPT) de diseño:", min_value=1, max_value=100, value=15, step=1)
    
    st.markdown("---")
    st.header("2. 🏜️ Tipos de Suelo")
    tipo_suelo = st.selectbox(
        "Tipo de Suelo (Aplicación):",
        ["Arenas", "Gravas", "Limos", "Suelos Intermedios", "Mostrar Todo"],
        index=0
    )
    st.caption(f"Filtra los métodos aplicables a **{tipo_suelo}**.")

# --- LÓGICA ---
df_completo = calcular_datos_base(n_spt)
df_filtrado = aplicar_filtro_suelo(df_completo, tipo_suelo)

cambio_filtro = (tipo_suelo != st.session_state.tipo_suelo_previo)
cambio_n = (n_spt != st.session_state.n_previo)

if st.session_state.tipo_suelo_previo is None or cambio_filtro:
    st.session_state.selecciones = {
        f"{row['Autor']}_{row['Aplicación']}": True 
        for _, row in df_filtrado.iterrows()
    }
    st.session_state.tipo_suelo_previo = tipo_suelo

if cambio_n:
    st.session_state.n_previo = n_spt

def get_seleccion(row):
    key = f"{row['Autor']}_{row['Aplicación']}"
    return st.session_state.selecciones.get(key, True)

df_filtrado.insert(0, "Seleccionar", df_filtrado.apply(get_seleccion, axis=1))

# --- EDITOR DE DATOS ---
st.subheader(f"1.🏜️Tipo de Suelo: {tipo_suelo}")

col_config = {
    "Seleccionar": st.column_config.CheckboxColumn("Incluir", width="small"),
    "E (MPa)": st.column_config.NumberColumn("E (MPa)", format="%.2f")
}

df_editado = st.data_editor(
    df_filtrado,
    column_config=col_config,
    disabled=["Autor", "Aplicación", "Fórmula Original", "E (MPa)"],
    hide_index=True,
    use_container_width=True,
    key="editor_principal"
)

for _, row in df_editado.iterrows():
    key = f"{row['Autor']}_{row['Aplicación']}"
    st.session_state.selecciones[key] = row["Seleccionar"]

# --- RESULTADOS ---
df_final = df_editado[df_editado["Seleccionar"] == True].drop(columns=["Seleccionar"])
num_seleccionados = len(df_final)

if num_seleccionados > 0:
    st.subheader("2.📊 Visualización Gráfica")
    
    # Texto combinado
    df_final["Texto_Barra"] = df_final["Autor"] + " (" + df_final["Aplicación"] + "): " + df_final["E (MPa)"].map('{:.1f}'.format) + " MPa"
    
    # Índice único para separar barras
    df_grafico = df_final.reset_index(drop=True)
    df_grafico["Indice"] = df_grafico.index.astype(str)

    fig = px.bar(
        df_grafico, 
        x="E (MPa)", 
        y="Indice", # Eje categórico forzado para separar
        text="Texto_Barra",
        color="Aplicación", 
        orientation='h', 
        title=f"Módulo de Elasticidad (N={n_spt}) - {tipo_suelo}", 
        color_discrete_sequence=px.colors.qualitative.Prism
    )
    
    # --- CONFIGURACIÓN PARA IGUALAR TAMAÑOS ---
    fig.update_layout(
        uniformtext_minsize=14, 
        uniformtext_mode='show',
        yaxis={'visible': False, 'showticklabels': False},
        xaxis_title="E (MPa)",
        legend=dict(
            orientation="h",
            yanchor="top",       # Anclado a la parte superior de la leyenda
            y=-0.15,             # Colocado debajo del eje X
            xanchor="center",    # Centrado horizontalmente
            x=0.5
        ),
        margin=dict(l=0, r=0, t=40, b=100), # Margen inferior aumentado para la leyenda
        height=200 + (len(df_final) * 50) 
    )
    
    fig.update_traces(
        textposition='inside', 
        insidetextanchor='start',
        textfont_size=14,     # Tamaño fijo base
        textfont_color='black' # Contraste garantizado (dentro/fuera)
    )
    
    st.plotly_chart(fig, use_container_width=True)
    
    st.subheader("3.💻Análisis Estadístico")
    df_stats = None
    if num_seleccionados >= 2:
        df_stats = pd.DataFrame({
            "Mínimo": [df_final["E (MPa)"].min()], 
            "Máximo": [df_final["E (MPa)"].max()],
            "Promedio": [df_final["E (MPa)"].mean()], 
            "Mediana": [df_final["E (MPa)"].median()],
            "Desv. Típica": [df_final["E (MPa)"].std()]
        })
        
        st.dataframe(
            df_stats.style.format("{:.2f}")
            .set_table_styles([
                {'selector': 'th', 'props': [('text-align', 'right')]},
                {'selector': 'td', 'props': [('text-align', 'right')]}
            ]), 
            hide_index=True, 
            use_container_width=True
        )
    else:
        st.warning("⚠️ Seleccione al menos 2 métodos para ver estadísticas.")

    st.markdown("---")
    with st.expander("📚 Ver Referencias Bibliográficas"):
        for ref in BIBLIOGRAFIA:
            st.markdown(f"- {ref}")

    st.subheader("4.📜Generar Informe")
    docx_file = generar_docx(n_spt, df_final, df_stats, fig, tipo_suelo)
    st.download_button("📄 Descargar Informe Word (.docx)", 
    docx_file, f"Informe_E_granular.docx", 
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",type="primary")

else:
    st.warning("⚠️ No hay métodos seleccionados para este grupo.")

# Disclaimer
st.info("⚠️ **Aviso de Responsabilidad:** Estas correlaciones son empíricas. Se recomienda contrastar estos valores con ensayos in situ (presiómetro, dilatómetro) o de laboratorio.")