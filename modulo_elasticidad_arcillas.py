import streamlit as st
import pandas as pd
import plotly.express as px
from io import BytesIO
import copy
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

# ==============================================================================
# CONFIGURACIÓN Y CONSTANTES
# ==============================================================================
st.set_page_config(page_title="E Suelos Cohesivos (Arcillas)", layout="wide")

BIBLIOGRAFIA = [
    "CTE DB SE-C. Código Técnico de la Edificación. Seguridad Estructural - Cimientos.",
    "Stroud, M. A. (1974). The standard penetration test in insensitive clays and soft rocks. Proceedings of the ESOPT, Stockholm.",
    "Butler, F. G. (1975). Heavily overconsolidated clays. Settlement of Structures.",
    "Bowles, J. E. (1996). Foundation Analysis and Design (5th ed.). McGraw-Hill."
]

# ==========================================
# 1. LÓGICA MATEMÁTICA
# ==========================================
def obtener_factor_f2_stroud(ip_val):
    """
    Obtiene el factor f2 (MPa/N) interpolando la gráfica de Stroud (1974)
    basada en el Índice de Plasticidad (IP).
    Valores aproximados de la curva estándar.
    """
    # Puntos de la curva (IP, f2)
    # f2 disminuye a medida que IP aumenta
    if ip_val <= 10: return 2.0
    if ip_val >= 60: return 0.6
    
    # Interpolación lineal simple entre puntos clave
    # (10, 2.0), (20, 1.7), (30, 1.4), (40, 1.0), (50, 0.8), (60, 0.6)
    puntos = {10: 2.0, 20: 1.7, 30: 1.4, 40: 1.0, 50: 0.8, 60: 0.6}
    
    # Buscar tramo
    x_keys = sorted(puntos.keys())
    for i in range(len(x_keys)-1):
        x1, x2 = x_keys[i], x_keys[i+1]
        if x1 <= ip_val <= x2:
            y1, y2 = puntos[x1], puntos[x2]
            return y1 + (ip_val - x1) * (y2 - y1) / (x2 - x1)
    return 0.6

def calcular_datos_arcillas(N_spt: int, IP: float, Cu_kPa: float, OCR_cat: str):
    """
    Calcula E para arcillas usando Stroud y CTE.
    """
    resultados = []
    
    # --- 1. STROUD (1974) ---
    # E = f2 * N_spt (MPa)
    f2 = obtener_factor_f2_stroud(IP)
    
    # Stroud suele dar un rango. Asumiremos el calculado como 'Valor Medio'
    # y crearemos un rango +/- 20% para simular Límite Superior/Inferior del Excel
    # o usamos la correlación directa.
    
    val_stroud = f2 * N_spt
    
    # Replicando lógica aproximada del Excel (Upper/Lower)
    resultados.append({
        "Autor": "Stroud (1974)", 
        "Aplicación": f"Límite Superior (IP={IP})", 
        "Fórmula Original": f"f2(IP)·N ({f2*1.2:.2f}·N)", 
        "E (MPa)": val_stroud * 1.2 # Asumiendo un rango superior
    })
    
    resultados.append({
        "Autor": "Stroud (1974)", 
        "Aplicación": f"Límite Inferior (IP={IP})", 
        "Fórmula Original": f"f2(IP)·N ({f2*0.8:.2f}·N)", 
        "E (MPa)": val_stroud * 0.8 # Asumiendo un rango inferior
    })

    # --- 2. STROUD Y BUTLER ---
    # Estos autores dan valores para 'Media' y 'Baja' plasticidad directos con N
    # Media Plasticidad: E ~ 0.5 N (MPa) ? (Revisando Excel: 14.7 para N=30 -> ~0.5)
    resultados.append({
        "Autor": "Stroud & Butler", 
        "Aplicación": "Arcillas Media Plasticidad", 
        "Fórmula Original": "0.5 · N (MPa)", 
        "E (MPa)": 0.5 * N_spt
    })
    
    # Baja Plasticidad: E ~ 0.6 N (MPa) ? (Revisando Excel: 17.6 para N=30 -> ~0.6)
    resultados.append({
        "Autor": "Stroud & Butler", 
        "Aplicación": "Arcillas Baja Plasticidad", 
        "Fórmula Original": "0.6 · N (MPa)", 
        "E (MPa)": 0.6 * N_spt
    })

    # --- 3. CTE DB SE-C (TABLA F.2) ---
    # E = K * Cu
    # K depende de IP y OCR.
    
    # Definir K según tabla (Valores del Excel como referencia)
    k_val = None
    condicion_txt = ""
    
    if IP < 30:
        if OCR_cat == "OCR < 3": k_val = 160; condicion_txt = "IP<30, OCR<3"
        elif OCR_cat == "3 < OCR < 5": k_val = 120; condicion_txt = "IP<30, 3<OCR<5"
        elif OCR_cat == "OCR > 5": k_val = 60; condicion_txt = "IP<30, OCR>5"
    elif 30 <= IP <= 50:
        if OCR_cat == "OCR < 3": k_val = 70; condicion_txt = "30<IP<50, OCR<3"
        elif OCR_cat == "3 < OCR < 5": k_val = 50; condicion_txt = "30<IP<50, 3<OCR<5"
        elif OCR_cat == "OCR > 5": k_val = 26; condicion_txt = "30<IP<50, OCR>5"
    else: # IP > 50
        if OCR_cat == "OCR < 3": k_val = 30; condicion_txt = "IP>50, OCR<3"
        elif OCR_cat == "3 < OCR < 5": k_val = 20; condicion_txt = "IP>50, 3<OCR<5"
        elif OCR_cat == "OCR > 5": k_val = 10; condicion_txt = "IP>50, OCR>5"

    if k_val:
        # Cu en kPa -> MPa es Cu/1000. Formula E = K * Cu_MPa
        # Pero los K del CTE (ej. 600) son para Cu en mismas unidades.
        # Si el Excel dice K=160, para Cu=200 kPa -> E = 160*200 = 32000 kPa = 32 MPa.
        # Si usamos Cu en MPa (0.2), E = 160 * 0.2 = 32 MPa.
        # Correcto.
        
        e_cte = k_val * (Cu_kPa / 1000.0) # Convertimos Cu a MPa primero
        
        resultados.append({
            "Autor": "CTE DB SE-C", 
            "Aplicación": f"Tabla F.2 ({condicion_txt})", 
            "Fórmula Original": f"{k_val} · Cu", 
            "E (MPa)": e_cte
        })
    else:
        # Fallback si algo falla
        resultados.append({
            "Autor": "CTE DB SE-C", 
            "Aplicación": "Fuera de rango", 
            "Fórmula Original": "-", 
            "E (MPa)": 0.0
        })

    return pd.DataFrame(resultados)

# ==========================================
# 2. GENERADOR DE INFORME WORD (Estilo CTE_2219)
# ==========================================
def generar_docx(n_val, ip_val, cu_val, ocr_val, df_final, df_stats, fig_plotly_original):
    doc = Document()
    
    # --- ESTILOS ---
    COLOR_TITULO = RGBColor(0x17, 0x36, 0x5D)
    COLOR_HEADING = RGBColor(0x36, 0x5F, 0x91)
    
    for section in doc.sections:
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)

    style_title = doc.styles['Title']
    style_title.font.name = 'Calibri Light'
    style_title.font.size = Pt(26)
    style_title.font.color.rgb = COLOR_TITULO
    
    style_h1 = doc.styles['Heading 1']
    style_h1.font.name = 'Calibri Light'
    style_h1.font.size = Pt(14)
    style_h1.font.color.rgb = COLOR_HEADING
    
    # --- CONTENIDO ---
    doc.add_heading('Informe estimación Módulo de Elasticidad en Arcillas', 0).alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    p_fecha = doc.add_paragraph()
    p_fecha.add_run(f'Fecha de emisión: {pd.Timestamp.now().strftime("%d/%m/%Y")}').bold = True
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Sección 1
    doc.add_heading('1. Datos de Entrada', level=1)
    
    # Lista de datos de entrada
    datos_entrada = [
        f"Valor N (SPT) de diseño: {n_val} golpes/30 cm",
        f"Índice de Plasticidad (IP): {ip_val} %",
        f"Cohesión sin drenaje (Cu): {cu_val} kPa",
        f"Grado de Sobreconsolidación: {ocr_val}"
    ]
    
    for dato in datos_entrada:
        p = doc.add_paragraph(dato, style='List Bullet')
        p.style.font.name = 'Calibri'

    # Sección 2 - Tabla
    doc.add_heading('2. Métodos de Cálculo Seleccionados', level=1)
    
    table = doc.add_table(rows=1, cols=4)
    try:
        table.style = 'Light List Accent 1'
    except:
        table.style = 'Table Grid'
        
    table.autofit = False 
    table.allow_autofit = False
    
    widths = [Inches(2.0), Inches(2.0), Inches(1.5), Inches(1.0)]
    headers = ['Autor', 'Aplicación', 'Fórmula', 'E (MPa)']
    
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.width = widths[i]
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
    for _, row in df_final.iterrows():
        row_cells = table.add_row().cells
        textos = [str(row['Autor']), str(row['Aplicación']), str(row['Fórmula Original']), f"{row['E (MPa)']:.2f}"]
        for idx, txt in enumerate(textos):
            row_cells[idx].text = txt
            row_cells[idx].width = widths[idx]
            row_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            row_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sección 3 - Estadísticas
    doc.add_heading('3. Análisis Estadístico', level=1)
    if df_stats is not None:
        doc.add_paragraph('Resumen estadístico de los métodos seleccionados:')
        stat_table = doc.add_table(rows=1, cols=5)
        try:
            stat_table.style = 'Light List Accent 1'
        except:
            stat_table.style = 'Table Grid'
        stat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        stat_table.autofit = True
        
        headers_stat = ['Mínimo', 'Máximo', 'Promedio', 'Mediana', 'Desv. Típica']
        for i, h in enumerate(headers_stat):
            cell = stat_table.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT 
            cell.paragraphs[0].runs[0].bold = True

        vals = stat_table.add_row().cells
        data_vals = [f"{df_stats.iloc[0][k]:.2f}" for k in headers_stat]
        for i, val in enumerate(data_vals):
            vals[i].text = val
            vals[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT 

    # Sección 4 - Gráfica
    doc.add_heading('4. Gráfica Comparativa', level=1)
    try:
        fig_word = copy.deepcopy(fig_plotly_original)
        fig_word.update_traces(textfont_size=14, textfont_color='black')
        fig_word.update_layout(uniformtext_minsize=14, uniformtext_mode='show')
        img_bytes = fig_word.to_image(format="png", width=1300, height=len(df_final)*60 + 200, scale=3)
        doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.paragraphs[-1].add_run().add_picture(BytesIO(img_bytes), width=Inches(6.0))
    except Exception as e:
        doc.add_paragraph(f"[Gráfica no disponible: {str(e)}]")

    doc.add_heading('5. Referencias Bibliográficas', level=1)
    for ref in BIBLIOGRAFIA:
        p = doc.add_paragraph(ref, style='List Bullet')
        p.style.font.name = 'Calibri'
        p.paragraph_format.space_after = Pt(6) 

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# 3. INTERFAZ STREAMLIT
# ==========================================

st.title("🧱 Estimación Módulo de Elasticidad en Arcillas")
st.markdown("---")

# --- SIDEBAR (INPUTS) ---
with st.sidebar:
    st.header("1.📝 Parámetros del Suelo")
    
    n_spt = st.number_input("Valor N (SPT):", min_value=1, value=15, step=1, help="Golpeo N30 estándar")
    ip_val = st.number_input("Índice de Plasticidad (IP %):", min_value=0.0, value=20.0, step=1.0)
    cu_val = st.number_input("Cohesión sin drenaje (Cu kPa):", min_value=0.0, value=100.0, step=10.0)
    
    st.markdown("---")
    st.header("2. ⚖️ Estado del Suelo")
    ocr_val = st.selectbox(
        "Grado de Sobreconsolidación (OCR):",
        ["OCR < 3", "3 < OCR < 5", "OCR > 5"],
        index=0,
        help="Necesario para correlaciones del CTE"
    )

# --- CONTROL DE ESTADO ---
if 'selecciones_arc' not in st.session_state: st.session_state.selecciones_arc = {}
# Resetear selecciones si cambian inputs críticos no es estrictamente necesario aquí si queremos persistencia,
# pero ayuda a refrescar. Lo mantenemos simple.

# --- CÁLCULO ---
df_completo = calcular_datos_arcillas(n_spt, ip_val, cu_val, ocr_val)

# Inicializar selección
for _, row in df_completo.iterrows():
    key = f"{row['Autor']}_{row['Aplicación']}"
    if key not in st.session_state.selecciones_arc:
        st.session_state.selecciones_arc[key] = True

def get_seleccion(row):
    key = f"{row['Autor']}_{row['Aplicación']}"
    return st.session_state.selecciones_arc.get(key, True)

df_completo.insert(0, "Seleccionar", df_completo.apply(get_seleccion, axis=1))

# --- EDITOR DE DATOS ---
st.subheader("1.📋 Tabla de Resultados")

col_config = {
    "Seleccionar": st.column_config.CheckboxColumn("Incluir", width="small"),
    "E (MPa)": st.column_config.NumberColumn("E (MPa)", format="%.2f")
}

df_editado = st.data_editor(
    df_completo,
    column_config=col_config,
    disabled=["Autor", "Aplicación", "Fórmula Original", "E (MPa)"],
    hide_index=True,
    use_container_width=True,
    key="editor_arcillas"
)

# Actualizar estado
for _, row in df_editado.iterrows():
    key = f"{row['Autor']}_{row['Aplicación']}"
    st.session_state.selecciones_arc[key] = row["Seleccionar"]

# --- FILTRADO FINAL ---
df_final = df_editado[df_editado["Seleccionar"] == True].drop(columns=["Seleccionar"])
num_seleccionados = len(df_final)

if num_seleccionados > 0:
    st.subheader("2.📊 Visualización Gráfica")
    
    # Texto combinado
    df_final["Texto_Barra"] = df_final["Autor"] + ": " + df_final["E (MPa)"].map('{:.1f}'.format) + " MPa"
    
    # Índice único
    df_grafico = df_final.reset_index(drop=True)
    df_grafico["Indice"] = df_grafico.index.astype(str)

    fig = px.bar(
        df_grafico, 
        x="E (MPa)", 
        y="Indice", 
        text="Texto_Barra",
        color="Aplicación", 
        orientation='h', 
        title=f"Módulo de Elasticidad (N={n_spt}, IP={ip_val}, Cu={cu_val})", 
        color_discrete_sequence=px.colors.qualitative.Pastel
    )
    
    fig.update_layout(
        uniformtext_minsize=14, 
        uniformtext_mode='show',
        yaxis={'visible': False, 'showticklabels': False},
        xaxis_title="E (MPa)",
        legend=dict(
            orientation="h",
            yanchor="top",
            y=-0.15,
            xanchor="center",
            x=0.5
        ),
        margin=dict(l=0, r=0, t=40, b=100),
        height=200 + (len(df_final) * 50) 
    )
    
    fig.update_traces(
        textposition='inside', 
        insidetextanchor='start',
        textfont_size=14,     
        textfont_color='black'
    )
    
    st.plotly_chart(fig, use_container_width=True)
    
    st.subheader("3.💻 Análisis Estadístico")
    if num_seleccionados >= 1: # Permitir stats con 1 solo dato, aunque sea trivial
        df_stats = pd.DataFrame({
            "Mínimo": [df_final["E (MPa)"].min()], 
            "Máximo": [df_final["E (MPa)"].max()],
            "Promedio": [df_final["E (MPa)"].mean()], 
            "Mediana": [df_final["E (MPa)"].median()],
            "Desv. Típica": [df_final["E (MPa)"].std() if len(df_final)>1 else 0]
        })
        
        st.dataframe(
            df_stats.style.format("{:.2f}")
            .set_table_styles([
                {'selector': 'th', 'props': [('text-align', 'right')]},
                {'selector': 'td', 'props': [('text-align', 'right')]}
            ]), 
            hide_index=True, 
            use_container_width=True
        )

    st.markdown("---")
    st.subheader("4.📜 Generar Informe")
    docx_file = generar_docx(n_spt, ip_val, cu_val, ocr_val, df_final, df_stats, fig)
    st.download_button("📄 Descargar Informe Word (.docx)", 
    docx_file, f"Informe_E_arcillas.docx", 
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary")

else:
    st.warning("⚠️ No hay métodos seleccionados.")

st.info("⚠️ **Nota:** Los cálculos se basan en correlaciones empíricas para arcillas (Stroud 1974, CTE DB SE-C).")