import streamlit as st
from docx import Document
import io
import pandas as pd

def calcular_Cv(H_dren=None, t=None, U=None):
    """
    Calcula el coeficiente de consolidación (Cv) según los datos disponibles.
    Solo aplica fórmulas para las que todos los parámetros necesarios estén disponibles.
    """
    resultados = {}
    formulas_usadas = {}

    def valor_aceptable(valor):
        return valor if valor is not None and valor > 0 else None

    H_dren = valor_aceptable(H_dren)
    t = valor_aceptable(t)
    U = valor_aceptable(U)

    if H_dren is not None and t is not None and U is not None:
        # Fórmula de Casagrande (logarítmica)
        T = 0.197 if U == 50 else (0.848 if U == 90 else None)
        if T is not None:
            Cv = (T * H_dren**2) / t
            resultados['Casagrande (logarítmica)'] = Cv if Cv > 0 else "Valor no aceptable para los datos introducidos"
            formulas_usadas['Casagrande (logarítmica)'] = {'formula': f'Cv = (T × H_dren²) / t, T = {T}', 'parametros': ['H_dren', 't', 'U']}

        # Fórmula de Taylor (raíz cuadrada)
        if U == 90:
            T = 0.848
            Cv = (T * H_dren**2) / t
            resultados['Taylor (raíz cuadrada)'] = Cv if Cv > 0 else "Valor no aceptable para los datos introducidos"
            formulas_usadas['Taylor (raíz cuadrada)'] = {'formula': 'Cv = (0.848 × H_dren²) / t', 'parametros': ['H_dren', 't']}

    return resultados, formulas_usadas

def generar_informe(H_dren, t, U, resultados, formulas_usadas):
    doc = Document()
    doc.add_heading('Informe de Cálculo del Coeficiente de Consolidación (Cv)', level=1)

    doc.add_heading('Datos Introducidos', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells

    hdr_cells[0].text = 'Parámetro'
    hdr_cells[1].text = 'Valor'
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    datos = []
    if H_dren is not None:
        datos.append(("Longitud de drenaje (H_dren, cm)", H_dren))
    if t is not None:
        datos.append(("Tiempo (t, min)", t))
    if U is not None:
        datos.append(("Grado de consolidación (U, %)", U))

    for parametro, valor in datos:
        row_cells = table.add_row().cells
        row_cells[0].text = parametro
        row_cells[1].text = str(valor)

    doc.add_heading('Resultados', level=2)
    doc.add_paragraph(f"Se calcularon {len(resultados)} correlaciones para Cv con los datos proporcionados.")

    for metodo, valor in resultados.items():
        doc.add_heading(metodo, level=3)
        formula_info = formulas_usadas[metodo]
        p = doc.add_paragraph()
        run = p.add_run(f"Fórmula: {formula_info['formula']}")
        run.bold = True
        doc.add_paragraph(f"Parámetros usados: {', '.join(formula_info['parametros'])}")
        if isinstance(valor, str):
            doc.add_paragraph(f"Resultado: {valor}")
        else:
            doc.add_paragraph(f"Resultado: Cv = {valor:.4f} cm²/min")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def main():
    st.set_page_config(layout="wide")
    st.title("Calculadora de Coeficiente de Consolidación (Cv)")
    st.markdown("Introduce los datos disponibles para calcular Cv según diferentes métodos.")

    if 'data' not in st.session_state:
        st.session_state.data = {
            "Símbolo": ["H_dren", "t", "U"],
            "Descripción del Parámetro": [
                "Longitud de drenaje", "Tiempo", "Grado de consolidación"
            ],
            "Valor": [None, None, None],
            "Unidad": ["cm", "min", "%"]
        }

    col1, col2 = st.columns([1, 1])

    with col1:
        df = pd.DataFrame(st.session_state.data)
        edited_df = st.data_editor(
            df,
            num_rows="fixed",
            hide_index=True,
            column_config={
                "Símbolo": st.column_config.TextColumn("Símbolo", disabled=True),
                "Descripción del Parámetro": st.column_config.TextColumn("Descripción del Parámetro", disabled=True),
                "Unidad": st.column_config.TextColumn("Unidad", disabled=True)
            },
            key="data_editor"
        )

        if st.button("Calcular Cv"):
            try:
                H_dren = float(edited_df.loc[edited_df["Símbolo"] == "H_dren", "Valor"].values[0]) if edited_df.loc[edited_df["Símbolo"] == "H_dren", "Valor"].values[0] and float(edited_df.loc[edited_df["Símbolo"] == "H_dren", "Valor"].values[0]) > 0 else None
            except (ValueError, TypeError):
                H_dren = None

            try:
                t = float(edited_df.loc[edited_df["Símbolo"] == "t", "Valor"].values[0]) if edited_df.loc[edited_df["Símbolo"] == "t", "Valor"].values[0] and float(edited_df.loc[edited_df["Símbolo"] == "t", "Valor"].values[0]) > 0 else None
            except (ValueError, TypeError):
                t = None

            try:
                U = float(edited_df.loc[edited_df["Símbolo"] == "U", "Valor"].values[0]) if edited_df.loc[edited_df["Símbolo"] == "U", "Valor"].values[0] and float(edited_df.loc[edited_df["Símbolo"] == "U", "Valor"].values[0]) > 0 else None
            except (ValueError, TypeError):
                U = None

            resultados, formulas_usadas = calcular_Cv(H_dren=H_dren, t=t, U=U)

            if resultados:
                st.subheader("Resultados de Cv según métodos aplicables:")
                for metodo, valor in resultados.items():
                    st.markdown(f"**{metodo}**")
                    st.write(f"Fórmula: {formulas_usadas[metodo]['formula']}")
                    st.write(f"Parámetros usados: {', '.join(formulas_usadas[metodo]['parametros'])}")
                    if isinstance(valor, str):
                        st.write(f"Resultado: {valor}")
                    else:
                        st.write(f"Resultado: Cv = {valor:.4f} cm²/min")
                    st.markdown("---")

                informe_buffer = generar_informe(H_dren, t, U, resultados, formulas_usadas)
                st.download_button(
                    label="Descargar informe en Word",
                    data=informe_buffer,
                    file_name="informe_Cv.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.warning("No hay suficientes datos para aplicar ningún método.")

    with col2:
        with st.expander("Métodos Disponibles", expanded=False):
            st.markdown("""
            <style>
            .formula {
                border-left: 3px solid #4CAF50;
                padding-left: 10px;
                margin-bottom: 15px;
            }
            </style>
            """, unsafe_allow_html=True)

            metodos = [
                {"name": "Casagrande (logarítmica)", "formula": "Cv = (T × H_dren²) / t, T = 0.197 para U = 50% y T = 0.848 para U = 90%", "params": "H_dren, t, U"},
                {"name": "Taylor (raíz cuadrada)", "formula": "Cv = (0.848 × H_dren²) / t, para U = 90%", "params": "H_dren, t"}
            ]

            for metodo in metodos:
                st.markdown(f'<div class="formula"><b>{metodo["name"]}</b><br>Fórmula: {metodo["formula"]}<br>Parámetros: {metodo["params"]}</div>', unsafe_allow_html=True)

if __name__ == "__main__":
    main()
