import streamlit as st
from docx import Document
import io
import pandas as pd

def calcular_Cc(LL=None, PL=None, IP=None, w=None, e=None, Gs=None, F=None):
    """
    Calcula el índice de compresión (Cc) según los datos disponibles.
    Solo aplica fórmulas para las que todos los parámetros necesarios estén disponibles.
    """
    resultados = {}
    formulas_usadas = {}

    def valor_aceptable(valor):
        return valor if valor is not None and valor >= 0 else "Valor no aceptable para los datos introducidos"

    # Fórmulas que solo requieren LL
    if LL is not None:
        cc_value = 0.009 * (LL - 10)
        resultados['Terzaghi & Peck (1967)'] = valor_aceptable(cc_value)
        formulas_usadas['Terzaghi & Peck (1967)'] = {'formula': 'Cc = 0.009 × (LL - 10)', 'parametros': ['LL']}

        cc_value = 0.007 * (LL - 7)
        resultados['Azzouz et al. (1976, arcillas remoldeadas)'] = valor_aceptable(cc_value)
        formulas_usadas['Azzouz et al. (1976, arcillas remoldeadas)'] = {'formula': 'Cc = 0.007 × (LL - 7)', 'parametros': ['LL']}

        cc_value = 0.0046 * (LL - 9)
        resultados['Azzouz et al. (1976, arcillas brasileñas)'] = valor_aceptable(cc_value)
        formulas_usadas['Azzouz et al. (1976, arcillas brasileñas)'] = {'formula': 'Cc = 0.0046 × (LL - 9)', 'parametros': ['LL']}

        cc_value = (LL - 13) / 109
        resultados['Mayne (1980)'] = valor_aceptable(cc_value)
        formulas_usadas['Mayne (1980)'] = {'formula': 'Cc = (LL - 13) / 109', 'parametros': ['LL']}

    # Fórmulas que requieren e
    if e is not None:
        cc_value = 0.3 * (e - 0.27)
        resultados['Hough (1957)'] = valor_aceptable(cc_value)
        formulas_usadas['Hough (1957)'] = {'formula': 'Cc = 0.3 × (e - 0.27)', 'parametros': ['e']}

        cc_value = 0.156 * e + 0.0107
        resultados['Azzouz et al. (1976, todas las arcillas)'] = valor_aceptable(cc_value)
        formulas_usadas['Azzouz et al. (1976, todas las arcillas)'] = {'formula': 'Cc = 0.156 × e + 0.0107', 'parametros': ['e']}

        cc_value = 0.75 * (e - 0.5)
        resultados['Azzouz et al. (1976, baja plasticidad)'] = valor_aceptable(cc_value)
        formulas_usadas['Azzouz et al. (1976, baja plasticidad)'] = {'formula': 'Cc = 0.75 × (e - 0.5)', 'parametros': ['e']}

        cc_value = 1.21 + 1.005 * (e - 1.87)
        resultados['Azzouz et al. (1976, São Paulo)'] = valor_aceptable(cc_value)
        formulas_usadas['Azzouz et al. (1976, São Paulo)'] = {'formula': 'Cc = 1.21 + 1.005 × (e - 1.87)', 'parametros': ['e']}

        cc_value = 1.15 * (e - 0.35)
        resultados['Nishida (1956)'] = valor_aceptable(cc_value)
        formulas_usadas['Nishida (1956)'] = {'formula': 'Cc = 1.15 × (e - 0.35)', 'parametros': ['e']}

    # Fórmulas que requieren w
    if w is not None:
        cc_value = 0.0115 * w
        resultados['Azzouz et al. (1976, suelos orgánicos)'] = valor_aceptable(cc_value)
        formulas_usadas['Azzouz et al. (1976, suelos orgánicos)'] = {'formula': 'Cc = 0.0115 × w', 'parametros': ['w']}

        cc_value = 0.0093 * w
        resultados['Koppula (1981, a)'] = valor_aceptable(cc_value)
        formulas_usadas['Koppula (1981, a)'] = {'formula': 'Cc = 0.0093 × w', 'parametros': ['w']}

        cc_value = 17.66e-5 * w**2 + 5.93e-3 * w - 0.135
        resultados['Azzouz et al. (1976, Chicago 2)'] = valor_aceptable(cc_value)
        formulas_usadas['Azzouz et al. (1976, Chicago 2)'] = {'formula': 'Cc = 17.66 × 10⁻⁵ × w² + 5.93 × 10⁻³ × w - 0.135', 'parametros': ['w']}

    # Fórmulas que requieren e, LL, y w
    if e is not None and LL is not None and w is not None:
        cc_value = 0.37 * (e + 0.003 * LL + 0.0004 * w - 0.34)
        resultados['Azzouz et al. (1976, 678 datos)'] = valor_aceptable(cc_value)
        formulas_usadas['Azzouz et al. (1976, 678 datos)'] = {'formula': 'Cc = 0.37 × (e + 0.003 × LL + 0.0004 × w - 0.34)', 'parametros': ['e', 'LL', 'w']}

    # Fórmulas que requieren PL y Gs (corregida según la imagen)
    if PL is not None and Gs is not None:
        cc_value = 0.005 * Gs * IP
        resultados['Wroth & Wood (1978)'] = valor_aceptable(cc_value)
        formulas_usadas['Wroth & Wood (1978)'] = {'formula': 'Cc = 0.005 × Gs × IP', 'parametros': ['Gs', 'IP']}

    # Fórmulas que requieren IP
    if IP is not None:
        cc_value = 0.046 + 0.0104 * IP
        resultados['Nakase et al. (1988)'] = valor_aceptable(cc_value)
        formulas_usadas['Nakase et al. (1988)'] = {'formula': 'Cc = 0.046 + 0.0104 × IP', 'parametros': ['IP']}

    # Fórmulas que requieren LL, IP, w, e, y F
    if LL is not None and IP is not None and w is not None and e is not None and F is not None:
        cc_value = -0.0997 + 0.009 * LL + 0.0014 * IP + 0.0036 * w + 0.1156 * e + 0.0025 * F
        resultados['Koppula (1981, b)'] = valor_aceptable(cc_value)
        formulas_usadas['Koppula (1981, b)'] = {
            'formula': 'Cc = -0.0997 + 0.009 × LL + 0.0014 × IP + 0.0036 × w + 0.1156 × e + 0.0025 × F',
            'parametros': ['LL', 'IP', 'w', 'e', 'F']
        }

    # Fórmulas que requieren w, Gs, PL, IP, y F
    if w is not None and Gs is not None and PL is not None and IP is not None and F is not None:
        cc_value = 0.329 * (0.01 * w * Gs - 0.027 * PL + 0.0133 * IP * (1.192 + F / IP))
        resultados['Carrier (1985)'] = valor_aceptable(cc_value)
        formulas_usadas['Carrier (1985)'] = {
            'formula': 'Cc = 0.329 × (0.01 × w × Gs - 0.027 × PL + 0.0133 × IP × (1.192 + F / IP))',
            'parametros': ['w', 'Gs', 'PL', 'IP', 'F']
        }

    return resultados, formulas_usadas

def generar_informe(LL, PL, IP, w, e, Gs, F, resultados, formulas_usadas):
    doc = Document()
    doc.add_heading('Informe de Cálculo del Índice de Compresión (Cc)', level=1)

    # Datos introducidos
    doc.add_heading('Datos Introducidos', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells

    # Encabezados en negrita
    hdr_cells[0].text = 'Parámetro'
    hdr_cells[1].text = 'Valor'
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Añadir los datos introducidos
    datos = []
    if LL is not None:
        datos.append(("Límite líquido (LL, %)", LL))
    if PL is not None:
        datos.append(("Límite plástico (PL, %)", PL))
    if IP is not None:
        datos.append(("Índice plástico (IP, %)", IP))
    if w is not None:
        datos.append(("Contenido de humedad (w, %)", w))
    if e is not None:
        datos.append(("Relación de vacíos (e)", e))
    if Gs is not None:
        datos.append(("Gravedad específica (Gs)", Gs))
    if F is not None:
        datos.append(("Porcentaje de grano fino (F, %)", F))

    for parametro, valor in datos:
        row_cells = table.add_row().cells
        row_cells[0].text = parametro
        row_cells[1].text = str(valor)

    # Resultados
    doc.add_heading('Resultados', level=2)
    doc.add_paragraph(f"Se calcularon {len(resultados)} correlaciones para Cc con los datos proporcionados.")

    for metodo, valor in resultados.items():
        doc.add_heading(metodo, level=3)
        formula_info = formulas_usadas[metodo]
        p = doc.add_paragraph()
        run = p.add_run(f"Fórmula: {formula_info['formula']}")
        run.bold = True
        doc.add_paragraph(f"Parámetros usados: {', '.join(formula_info['parametros'])}")
        if isinstance(valor, str):
            doc.add_paragraph(f"Resultado: {valor}")
        else:
            doc.add_paragraph(f"Resultado: Cc = {valor:.4f}")

    # Guardar el documento en un buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def main():
    st.set_page_config(layout="wide")
    st.title("Calculadora de Índice de Compresión (Cc)")
    st.markdown("Introduce los datos disponibles para calcular Cc según diferentes fórmulas.")

    # Crear columnas para la tabla y el panel lateral
    col1, col2 = st.columns([1, 1])

    # Tabla de entrada de datos (izquierda)
    with col1:
        data = {
            "Symbol": ["LL", "PL", "IP", "w", "e", "Gs", "gsat", "gdry", "F", "n"],
            "Parameter Description": [
                "Liquid Limit", "Plastic Limit", "Plastic Index", "Moisture Content",
                "Void Ratio", "Specific Gravity", "Saturated Unit Weight", "Dry Unit Weight",
                "Percent fine-grained (clay & silt)", "Porosity"
            ],
            "Value": [None, None, None, None, None, None, None, None, None, None],
            "Unit": ["%", "%", "%", "%", "", "", "kN/m³", "kN/m³", "%", ""]
        }

        df = pd.DataFrame(data)
        edited_df = st.data_editor(
            df,
            num_rows="fixed",
            hide_index=True,
            column_config={
                "Symbol": st.column_config.TextColumn("Symbol", disabled=True),
                "Parameter Description": st.column_config.TextColumn("Parameter Description", disabled=True),
                "Unit": st.column_config.TextColumn("Unit", disabled=True)
            }
        )

    # Panel lateral con las fórmulas (derecha)
    with col2:
        st.header("Fórmulas Disponibles")
        st.markdown("""
        <style>
        .formula {
            border-left: 3px solid #4CAF50;
            padding-left: 10px;
            margin-bottom: 15px;
        }
        </style>
        """, unsafe_allow_html=True)

        formulas = [
            {"name": "Terzaghi & Peck (1967)", "formula": "Cc = 0.009 × (LL - 10)", "params": "LL"},
            {"name": "Azzouz et al. (1976, arcillas remoldeadas)", "formula": "Cc = 0.007 × (LL - 7)", "params": "LL"},
            {"name": "Azzouz et al. (1976, arcillas brasileñas)", "formula": "Cc = 0.0046 × (LL - 9)", "params": "LL"},
            {"name": "Mayne (1980)", "formula": "Cc = (LL - 13) / 109", "params": "LL"},
            {"name": "Hough (1957)", "formula": "Cc = 0.3 × (e - 0.27)", "params": "e"},
            {"name": "Azzouz et al. (1976, todas las arcillas)", "formula": "Cc = 0.156 × e + 0.0107", "params": "e"},
            {"name": "Azzouz et al. (1976, baja plasticidad)", "formula": "Cc = 0.75 × (e - 0.5)", "params": "e"},
            {"name": "Azzouz et al. (1976, São Paulo)", "formula": "Cc = 1.21 + 1.005 × (e - 1.87)", "params": "e"},
            {"name": "Nishida (1956)", "formula": "Cc = 1.15 × (e - 0.35)", "params": "e"},
            {"name": "Azzouz et al. (1976, suelos orgánicos)", "formula": "Cc = 0.0115 × w", "params": "w"},
            {"name": "Koppula (1981, a)", "formula": "Cc = 0.0093 × w", "params": "w"},
            {"name": "Azzouz et al. (1976, Chicago 2)", "formula": "Cc = 17.66 × 10⁻⁵ × w² + 5.93 × 10⁻³ × w - 0.135", "params": "w"},
            {"name": "Azzouz et al. (1976, 678 datos)", "formula": "Cc = 0.37 × (e + 0.003 × LL + 0.0004 × w - 0.34)", "params": "e, LL, w"},
            {"name": "Wroth & Wood (1978)", "formula": "Cc = 0.005 × Gs × IP", "params": "Gs, IP"},
            {"name": "Nakase et al. (1988)", "formula": "Cc = 0.046 + 0.0104 × IP", "params": "IP"},
            {"name": "Koppula (1981, b)", "formula": "Cc = -0.0997 + 0.009 × LL + 0.0014 × IP + 0.0036 × w + 0.1156 × e + 0.0025 × F", "params": "LL, IP, w, e, F"},
            {"name": "Carrier (1985)", "formula": "Cc = 0.329 × (0.01 × w × Gs - 0.027 × PL + 0.0133 × IP × (1.192 + F / IP))", "params": "w, Gs, PL, IP, F"}
        ]

        for formula in formulas:
            st.markdown(f'<div class="formula"><b>{formula["name"]}</b><br>Fórmula: {formula["formula"]}<br>Parámetros: {formula["params"]}</div>', unsafe_allow_html=True)

    # Extraer los valores introducidos
    LL = edited_df.loc[edited_df["Symbol"] == "LL", "Value"].iloc[0]
    PL = edited_df.loc[edited_df["Symbol"] == "PL", "Value"].iloc[0]
    IP = edited_df.loc[edited_df["Symbol"] == "IP", "Value"].iloc[0]
    w = edited_df.loc[edited_df["Symbol"] == "w", "Value"].iloc[0]
    e = edited_df.loc[edited_df["Symbol"] == "e", "Value"].iloc[0]
    Gs = edited_df.loc[edited_df["Symbol"] == "Gs", "Value"].iloc[0]
    F = edited_df.loc[edited_df["Symbol"] == "F", "Value"].iloc[0]

    if st.button("Calcular Cc"):
        resultados, formulas_usadas = calcular_Cc(LL=LL, PL=PL, IP=IP, w=w, e=e, Gs=Gs, F=F)

        if resultados:
            st.subheader("Resultados de Cc según fórmulas aplicables:")
            for metodo, valor in resultados.items():
                st.markdown(f"**{metodo}**")
                st.write(f"Fórmula: {formulas_usadas[metodo]['formula']}")
                st.write(f"Parámetros usados: {', '.join(formulas_usadas[metodo]['parametros'])}")
                if isinstance(valor, str):
                    st.write(f"Resultado: {valor}")
                else:
                    st.write(f"Resultado: Cc = {valor:.4f}")
                st.markdown("---")

            # Generar informe
            informe_buffer = generar_informe(LL, PL, IP, w, e, Gs, F, resultados, formulas_usadas)
            st.download_button(
                label="Descargar informe en Word",
                data=informe_buffer,
                file_name="informe_Cc.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.warning("No hay suficientes datos para aplicar ninguna fórmula.")

if __name__ == "__main__":
    main()
